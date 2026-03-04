import os
import io
import pandas as pd
from docx import Document
import zipfile
from parsers import parse_document
from database import init_db
from app import app
from categorizer import categorize_transaction

os.environ['TESTING'] = 'true'
init_db()

# 1. XLSX Validation
df = pd.DataFrame({'Date': ['02/15/2026'], 'Description': ['UBER *TRIP'], 'Amount': [-25.50]})
df.to_excel('test_verify.xlsx', index=False)

t_xlsx, ai_xlsx = parse_document('test_verify.xlsx', 'auto')
print(f'XLSX extracted: {len(t_xlsx)} records')
assert t_xlsx[0]['amount'] == -25.5
assert t_xlsx[0]['description'] == 'UBER *TRIP'

cat1 = categorize_transaction(1, t_xlsx[0]['description'], t_xlsx[0]['amount'], 'debit', '')
print(f'XLSX Categorized as: {cat1["category"]} (Personal: {cat1["is_personal"]})')

# 2. DOCX Extraction Validation
doc = Document()
doc.add_paragraph('This is a test paragraph for Forensic evidence.')
table = doc.add_table(rows=1, cols=2)
cells = table.rows[0].cells
cells[0].text = 'Header 1'
cells[1].text = 'Header 2'
doc.save('test_verify.docx')

t_docx, ai_docx = parse_document('test_verify.docx', 'proof')
print(f'DOCX text blocks extracted: {len(ai_docx["content"]["text"])} paragraphs')
text_blob = " ".join(ai_docx['content']['text']) + " " + str(ai_docx['content']['tables'])
assert 'Forensic evidence' in text_blob
assert 'Header 1' in text_blob

os.remove('test_verify.xlsx')
os.remove('test_verify.docx')
print('Extraction and Categorization tests passed safely.')
