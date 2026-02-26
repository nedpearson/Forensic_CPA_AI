"""
Document parsers for bank statements, credit card statements, Venmo exports, Excel, and Word.
Extracts transactions from uploaded PDFs and spreadsheets.
"""
import re
import os
import hashlib
import pdfplumber
import pandas as pd
from datetime import datetime
from openpyxl import load_workbook
from docx import Document as DocxDocument
from pydantic import BaseModel, Field
from typing import List, Optional
import json
TESSERACT_CMD = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\Users\nedpe\AppData\Local\poppler\poppler-24.08.0\Library\bin'


def parse_pdf_text(filepath):
    """Extract all text from a PDF file, page by page. Falls back to OCR for scanned/image PDFs."""
    pages = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and len(text.strip()) > 20:
                pages.append(text)

    # If pdfplumber got very little text, try OCR
    total_chars = sum(len(p) for p in pages)
    if total_chars < 100:
        ocr_pages, ocr_error = _ocr_pdf(filepath)
        if ocr_pages:
            return ocr_pages
        elif ocr_error:
            raise Exception(f"This PDF appears to be a scanned image, but OCR failed: {ocr_error}")

    return pages


def _ocr_pdf(filepath):
    """Use Tesseract OCR to extract text from image-based PDF pages."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

        if not os.path.exists(POPPLER_PATH):
            return [], "Poppler binary not found at " + POPPLER_PATH
        if not os.path.exists(TESSERACT_CMD):
            return [], "Tesseract binary not found at " + TESSERACT_CMD

        images = convert_from_path(filepath, dpi=300, poppler_path=POPPLER_PATH)
        pages = []
        for img in images:
            text = pytesseract.image_to_string(img)
            if text and text.strip():
                pages.append(text)
        return pages, None
    except ImportError:
        return [], "Python packages 'pytesseract' or 'pdf2image' are not installed."
    except Exception as e:
        return [], str(e)


def _normalize_text(text):
    if not text:
        return ""
    # trim, collapse whitespace, uppercase
    text = str(text)
    text = re.sub(r'\s+', ' ', text).strip().upper()
    return text

def compute_transaction_hash(account_scope_id, trans_date, amount, description, merchant=None, currency='USD', post_date=None, check_number=None, running_balance=None):
    """Create a deterministic hash for duplicate detection."""
    norm_desc = _normalize_text(description)
    norm_merch = _normalize_text(merchant) if merchant else ""
    
    # Amount in cents
    try:
        amount_cents = int(round(float(amount) * 100))
    except (ValueError, TypeError):
        amount_cents = 0
        
    components = [
        str(account_scope_id or ""),
        str(trans_date or ""),
        str(amount_cents),
        norm_desc,
        norm_merch,
        str(currency).upper(),
    ]
    if post_date:
        components.append(str(post_date))
    if check_number:
        components.append(str(check_number))
    if running_balance is not None:
        try:
            bal_cents = int(round(float(running_balance) * 100))
            components.append(str(bal_cents))
        except (ValueError, TypeError):
            pass
            
    key = "|".join(components)
    return hashlib.sha256(key.encode('utf-8')).hexdigest()


def parse_pdf_tables(filepath):
    """Extract tables from a PDF file."""
    all_tables = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for table in tables:
                all_tables.append({'page': i + 1, 'data': table})
    return all_tables


# --- Bank of St. Francisville Statement Parser ---

MONTHS = {'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04', 'may': '05', 'jun': '06',
          'jul': '07', 'aug': '08', 'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'}


def _normalize_bank_date(date_str, year='2026'):
    """Convert 'Jan 13' to '2026-01-13'."""
    parts = date_str.strip().split()
    if len(parts) == 2:
        mon = MONTHS.get(parts[0].lower()[:3], '')
        day = parts[1].zfill(2)
        if mon:
            return f"{year}-{mon}-{day}"
    return date_str


def parse_bank_statement(filepath):
    """Parse PDF statements using Generative AI with strict JSON schema."""
    pages = parse_pdf_text(filepath)
    full_text = "\n".join(pages)

    transactions = []
    account_info = {
        'institution': 'Unknown',
        'account_type': 'bank',
        'account_number': '',
        'account_name': '',
        'statement_start': '',
        'statement_end': '',
    }

    if not full_text.strip():
        return transactions, account_info

    try:
        from openai import OpenAI
        import httpx
        custom_http_client = httpx.Client()
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"), http_client=custom_http_client)
        
        class TransactionSchema(BaseModel):
            date: str = Field(description="ISO 8601 date, e.g. 2024-12-17")
            description: str = Field(description="Cleaned description of the transaction")
            amount: float = Field(description="Signed float. Negative for debits/withdrawals/fees, positive for credits/deposits.")
            type: str = Field(description="'debit' or 'credit'")
            balance: Optional[float] = Field(None, description="Running balance if shown")

        class AccountInfoSchema(BaseModel):
            institution: str = Field(description="Name of the bank or institution")
            account_type: str = Field(description="'bank', 'credit_card', or 'other'")
            account_number: str = Field(description="Account number or last 4 digits")
            statement_start: str = Field(description="Statement start date (YYYY-MM-DD), or empty if unknown")
            statement_end: str = Field(description="Statement end date (YYYY-MM-DD), or empty if unknown")
            
        class StatementExtractionSchema(BaseModel):
            account_info: AccountInfoSchema
            transactions: List[TransactionSchema]

        response = client.beta.chat.completions.parse(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a financial parsing assistant. Extract all transactions and account information from the given bank statement text. Ensure the amount is always negative for debits (withdrawals, fees) and positive for credits (deposits). Include all explicit line-item transactions, overdraft fees, service charges, or other fees that affect the balance."},
                {"role": "user", "content": full_text[:100000]} # Limit to 100k chars for safety
            ],
            response_format=StatementExtractionSchema,
            temperature=0.0
        )

        parsed_data = response.choices[0].message.parsed
        account_info = parsed_data.account_info.model_dump()

        # Map to old standard transactional format
        for trans in parsed_data.transactions:
            transactions.append({
                'trans_date': trans.date,
                'post_date': trans.date,
                'description': trans.description,
                'amount': trans.amount,
                'trans_type': trans.type,
                'balance': trans.balance,
                'cardholder_name': '',
                'card_last_four': '',
                'payment_method': ''
            })

    except Exception as e:
        print(f"LLM Parsing failed: {e}")
        # Could fallback to regex here if heavily needed...

    return transactions, account_info

    # Extract account number
    acct_match = re.search(r'Account\s*(?:Number|#|No\.?)[:\s]*(\d+)', full_text, re.IGNORECASE)
    if acct_match:
        account_info['account_number'] = acct_match.group(1)

    # Extract statement dates
    start_match = re.search(r'Statement\s+Date\s+(\d{2}/\d{2}/\d{4})', full_text)
    end_match = re.search(r'Statement\s+Thru\s+Date\s+(\d{2}/\d{2}/\d{4})', full_text)
    if start_match:
        account_info['statement_start'] = start_match.group(1)
    if end_match:
        account_info['statement_end'] = end_match.group(1)

    # Determine year from statement date
    year = '2026'
    if start_match:
        year = start_match.group(1)[-4:]

    # Also try old format
    if not start_match:
        date_match = re.search(r'(\w+\s+\d{1,2},?\s+\d{4})\s*(?:through|to|-)\s*(\w+\s+\d{1,2},?\s+\d{4})', full_text, re.IGNORECASE)
        if date_match:
            account_info['statement_start'] = date_match.group(1)
            account_info['statement_end'] = date_match.group(2)

    # Extract card last 4 from transaction descriptions
    card_numbers = set()

    # OCR format: "Jan DD <description> <amount> <balance>"
    # Multi-line: description may wrap, amounts are at end of line
    # Key patterns:
    #   Jan 13 POS PURCHASE NON-PIN VENMO *DAVID 500.00 5,214.69
    #   Jan 20 DEPOSIT 15,821.25 34,938.48
    #   Jan 20 DEGRAW CONSULTIN/SALE GULF COAST 14,000.00 33,182.47

    # Pattern: Month Day at start of line, then description, then 1-2 amounts at end
    line_pattern = re.compile(
        r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{1,2})\s+'  # date
        r'(.+?)\s+'                                                            # description
        r'([\d,]+\.\d{2})\s+'                                                 # amount1 (deposit or withdrawal)
        r'([\d,]+\.\d{2})\s*$',                                               # balance
        re.IGNORECASE
    )

    # Deposit pattern has amount in deposit column (before withdrawal column)
    # We'll determine deposit vs withdrawal by comparing to previous balance

    prev_balance = None
    in_transactions = False

    for page_text in pages:
        lines = page_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect transaction section
            if 'ACCOUNT TRANSACTIONS' in line.upper():
                in_transactions = True
                continue

            # Skip headers and footers
            skip_words = ['PAGE ', 'ACCOUNT SUMMARY', 'STATEMENT DATE', 'ACCOUNT NUMBER',
                          'DATE DESCRIPTION DEPOSITS', 'BEGINNING BALANCE', 'ENDING BALANCE',
                          'IN CASE OF ERRORS', 'MEMBER FDIC', 'BANK OF', 'ST. FRANCISVILL',
                          'SERVICE CHARGE', 'TOTAL SERVICE', 'CUSTOMER SERVICE',
                          'IF YOU THINK', 'DIRECT ALL', 'CHECK/ITEMS']
            if any(skip in line.upper() for skip in skip_words):
                continue

            # Skip OCR noise (very short lines, all caps headers, page markers)
            if len(line) < 10 or re.match(r'^[A-Z0-9\s]{5,20}$', line):
                continue

            match = line_pattern.match(line)
            if match:
                month_str = match.group(1)
                day = match.group(2)
                desc = match.group(3).strip()
                amount1 = float(match.group(4).replace(',', ''))
                balance = float(match.group(5).replace(',', ''))

                date_str = _normalize_bank_date(f"{month_str} {day}", year)

                # Extract card number from description (e.g., *****3992, *****8279)
                card_match = re.search(r'\*{3,5}(\d{4})', desc)
                card_last_four = card_match.group(1) if card_match else ''
                if card_last_four:
                    card_numbers.add(card_last_four)

                # Clean description: remove trailing date/time and card number
                clean_desc = re.sub(r'\s+\d{2}/\d{2}\s+\d{2}:\d{2}\s*$', '', desc)
                clean_desc = re.sub(r'\s+\*{3,5}\d{4}\s*', ' ', clean_desc).strip()
                clean_desc = re.sub(r'\s+', ' ', clean_desc)

                # Determine if deposit or withdrawal by balance change
                is_deposit = False
                if 'DEPOSIT' in desc.upper():
                    is_deposit = True
                elif prev_balance is not None:
                    # If balance went up, it's a deposit
                    is_deposit = balance > prev_balance
                else:
                    # No previous balance to compare - check column position heuristic
                    is_deposit = 'DEPOSIT' in desc.upper() or 'CREDIT' in desc.upper()

                if is_deposit:
                    trans_type = 'deposit'
                    amount = amount1
                else:
                    trans_type = 'debit'
                    amount = -amount1

                # Determine cardholder from card number
                cardholder = ''
                if card_last_four == '3992':
                    cardholder = 'JAMES HENDRICK'
                elif card_last_four == '8279':
                    cardholder = 'GERALD PEARSON'

                # Determine payment method
                payment_method = 'debit'
                if 'VENMO' in desc.upper():
                    payment_method = 'venmo'
                elif 'CHECK' in desc.upper():
                    payment_method = 'check'
                elif 'WIRE' in desc.upper() or 'ACH' in desc.upper():
                    payment_method = 'transfer'
                elif 'MOBILE' in desc.upper():
                    payment_method = 'mobile'

                transactions.append({
                    'trans_date': date_str,
                    'post_date': date_str,
                    'description': clean_desc,
                    'amount': amount,
                    'trans_type': trans_type,
                    'cardholder_name': cardholder,
                    'card_last_four': card_last_four,
                    'payment_method': payment_method,
                })

                prev_balance = balance

    return transactions, account_info


# --- Capital One Credit Card Statement Parser ---

def parse_capital_one_statement(filepath):
    """Parse Capital One Spark Cash Plus credit card PDF statements."""
    pages = parse_pdf_text(filepath)
    full_text = "\n".join(pages)

    transactions = []
    account_info = {
        'institution': 'Capital One',
        'account_type': 'credit_card',
        'account_number': '',
        'account_name': '',
        'statement_start': '',
        'statement_end': '',
    }

    # Extract account ending number
    acct_match = re.search(r'ending in (\d{4})', full_text)
    if acct_match:
        account_info['account_number'] = acct_match.group(1)

    # Extract billing cycle dates
    cycle_match = re.search(r'(\w{3}\s+\d{1,2},?\s+\d{4})\s*-\s*(\w{3}\s+\d{1,2},?\s+\d{4})', full_text)
    if cycle_match:
        account_info['statement_start'] = cycle_match.group(1)
        account_info['statement_end'] = cycle_match.group(2)

    # Extract cardholder name from header
    name_match = re.search(r'(GERALD T PEARSON|JAMES HENDRICK)', full_text)
    if name_match:
        account_info['account_name'] = name_match.group(1)

    # Parse transactions by cardholder section
    # Capital One statements group transactions under each cardholder
    current_cardholder = ''
    current_card = ''
    current_section = ''  # 'payments' or 'transactions'

    for page_text in pages:
        lines = page_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect cardholder section headers
            cardholder_match = re.match(r'([\w\s]+?)\s*#(\d{4}):\s*(Payments.*|Transactions)', line)
            if cardholder_match:
                current_cardholder = cardholder_match.group(1).strip()
                current_card = cardholder_match.group(2)
                section_type = cardholder_match.group(3)
                current_section = 'payment' if 'Payment' in section_type else 'transaction'
                continue

            # Skip headers
            if line.startswith('Trans Date') or line.startswith('Description'):
                continue

            # Parse transaction lines: "Mon DD  Mon DD  Description  $Amount"
            trans_match = re.match(
                r'(\w{3}\s+\d{1,2})\s+(\w{3}\s+\d{1,2})\s+(.+?)\s+\$?([\d,]+\.\d{2})\s*$',
                line
            )
            if not trans_match:
                # Try: "Mon DD  Mon DD  Description  -$Amount"
                trans_match = re.match(
                    r'(\w{3}\s+\d{1,2})\s+(\w{3}\s+\d{1,2})\s+(.+?)\s+-?\$?([\d,]+\.\d{2})\s*$',
                    line
                )

            if trans_match:
                trans_date = trans_match.group(1)
                post_date = trans_match.group(2)
                desc = trans_match.group(3).strip()
                amount_str = trans_match.group(4).replace(',', '')

                try:
                    amount = float(amount_str)
                except ValueError:
                    continue

                # Payments/credits are positive (reduce balance), charges are negative
                if current_section == 'payment':
                    trans_type = 'payment'
                    amount = amount  # credit
                else:
                    trans_type = 'debit'
                    amount = -amount  # charge

                # Check for fees
                if 'FEE' in desc.upper():
                    trans_type = 'fee'
                    amount = -abs(amount)

                transactions.append({
                    'trans_date': trans_date,
                    'post_date': post_date,
                    'description': desc,
                    'amount': amount,
                    'trans_type': trans_type,
                    'cardholder_name': current_cardholder,
                    'card_last_four': current_card,
                    'payment_method': 'credit',
                })

    # Also parse the Fees section at the bottom
    fees_section = re.search(r'Fees\s*\n(.*?)(?:Total Fees|$)', full_text, re.DOTALL)
    if fees_section:
        fee_lines = fees_section.group(1).strip().split('\n')
        for line in fee_lines:
            fee_match = re.match(
                r'(\w{3}\s+\d{1,2})\s+(\w{3}\s+\d{1,2})\s+(.+?)\s+\$?([\d,]+\.\d{2})',
                line.strip()
            )
            if fee_match:
                # Check if this fee was already added
                desc = fee_match.group(3).strip()
                amt = float(fee_match.group(4).replace(',', ''))
                already_exists = any(
                    t['description'] == desc and abs(t['amount']) == amt
                    for t in transactions
                )
                if not already_exists:
                    transactions.append({
                        'trans_date': fee_match.group(1),
                        'post_date': fee_match.group(2),
                        'description': desc,
                        'amount': -amt,
                        'trans_type': 'fee',
                        'cardholder_name': '',
                        'card_last_four': account_info.get('account_number', ''),
                        'payment_method': 'credit',
                    })

    return transactions, account_info


# --- Venmo Statement Parser ---

def parse_venmo_statement(filepath):
    """Parse Venmo CSV or PDF exports."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.csv':
        return parse_venmo_csv(filepath)
    elif ext == '.pdf':
        return parse_venmo_pdf(filepath)
    elif ext in ('.xlsx', '.xls'):
        return parse_venmo_excel(filepath)
    return [], {}


def parse_venmo_csv(filepath):
    """Parse Venmo CSV export."""
    df = pd.read_csv(filepath, skiprows=lambda x: x < 2 if x < 5 else False)
    transactions = []

    # Venmo CSV columns vary but typically include: ID, Datetime, Type, Status, Note, From, To, Amount
    for _, row in df.iterrows():
        try:
            amount_str = str(row.get('Amount (total)', row.get('Amount', '0')))
            amount_str = amount_str.replace('$', '').replace(',', '').replace('+', '').replace(' ', '')
            if not amount_str or amount_str == 'nan':
                continue
            amount = float(amount_str)

            desc_parts = []
            if pd.notna(row.get('Type')):
                desc_parts.append(str(row['Type']))
            if pd.notna(row.get('From')):
                desc_parts.append(f"From: {row['From']}")
            if pd.notna(row.get('To')):
                desc_parts.append(f"To: {row['To']}")
            if pd.notna(row.get('Note')):
                desc_parts.append(str(row['Note']))

            date_str = str(row.get('Datetime', row.get('Date', '')))

            transactions.append({
                'trans_date': date_str,
                'post_date': date_str,
                'description': ' | '.join(desc_parts) if desc_parts else 'Venmo Transaction',
                'amount': amount,
                'trans_type': 'credit' if amount > 0 else 'debit',
                'cardholder_name': str(row.get('From', '')) if amount > 0 else str(row.get('To', '')),
                'card_last_four': '',
                'payment_method': 'venmo',
            })
        except (ValueError, TypeError):
            continue

    account_info = {
        'institution': 'Venmo',
        'account_type': 'venmo',
        'account_number': 'venmo',
    }
    return transactions, account_info


def parse_venmo_excel(filepath):
    """Parse Venmo data from Excel export."""
    df = pd.read_excel(filepath)
    # Reuse CSV logic since structure is similar
    transactions = []
    for _, row in df.iterrows():
        try:
            amount = 0
            for col in df.columns:
                if 'amount' in col.lower():
                    val = str(row[col]).replace('$', '').replace(',', '').replace('+', '')
                    if val and val != 'nan':
                        amount = float(val)
                        break

            desc = ' | '.join(str(row[c]) for c in df.columns if pd.notna(row[c]) and 'amount' not in c.lower() and 'id' not in c.lower())

            transactions.append({
                'trans_date': str(row.get('Date', row.get('Datetime', ''))),
                'post_date': '',
                'description': desc or 'Venmo Transaction',
                'amount': amount,
                'trans_type': 'credit' if amount > 0 else 'debit',
                'cardholder_name': '',
                'card_last_four': '',
                'payment_method': 'venmo',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Venmo', 'account_type': 'venmo', 'account_number': 'venmo'}


def parse_venmo_pdf(filepath):
    """Parse Venmo PDF statement."""
    pages = parse_pdf_text(filepath)
    transactions = []

    for page_text in pages:
        lines = page_text.split('\n')
        for line in lines:
            # Venmo PDF patterns: date, name, +/- amount
            match = re.match(
                r'(\d{1,2}/\d{1,2}/\d{2,4})\s+(.+?)\s+([-+]?\$?[\d,]+\.\d{2})',
                line.strip()
            )
            if match:
                amount_str = match.group(3).replace('$', '').replace(',', '').replace('+', '')
                try:
                    amount = float(amount_str)
                except ValueError:
                    continue

                transactions.append({
                    'trans_date': match.group(1),
                    'post_date': match.group(1),
                    'description': f"Venmo: {match.group(2).strip()}",
                    'amount': amount,
                    'trans_type': 'credit' if amount > 0 else 'debit',
                    'cardholder_name': match.group(2).strip(),
                    'card_last_four': '',
                    'payment_method': 'venmo',
                })

    return transactions, {'institution': 'Venmo', 'account_type': 'venmo', 'account_number': 'venmo'}


# --- Excel Parser (generic) ---

def parse_excel_transactions(filepath):
    """Parse transactions from Excel files. Auto-detects column structure across all sheets."""
    all_dfs = pd.read_excel(filepath, sheet_name=None)
    transactions = []
    account_info = {
        'institution': 'Unknown',
        'account_type': 'bank',
        'account_number': '',
    }

    for sheet_name, df in all_dfs.items():
        if df.empty:
            continue

        # Clean column names to prevent KeyErrors from trailing whitespace
        df.columns = [str(c).strip() for c in df.columns]

        # Auto-detect column mappings
        col_map = {}
        for col in df.columns:
            col_lower = col.lower()
            if 'date' in col_lower and 'post' not in col_lower:
                col_map['trans_date'] = col
            elif 'post' in col_lower and 'date' in col_lower:
                col_map['post_date'] = col
            elif 'desc' in col_lower or 'memo' in col_lower or 'narrative' in col_lower:
                col_map['description'] = col
            elif 'amount' in col_lower or 'total' in col_lower:
                col_map['amount'] = col
            elif 'debit' in col_lower:
                col_map['debit'] = col
            elif 'credit' in col_lower:
                col_map['credit'] = col
            elif 'category' in col_lower or 'type' in col_lower:
                col_map['category'] = col
            elif 'card' in col_lower or 'holder' in col_lower or 'name' in col_lower:
                col_map['cardholder'] = col
            elif 'check' in col_lower and 'num' in col_lower:
                col_map['check_number'] = col

        for _, row in df.iterrows():
            try:
                # Get amount
                if 'amount' in col_map:
                    amount_val = row[col_map['amount']]
                    if pd.isna(amount_val):
                        continue
                    amount_str = str(amount_val).replace('$', '').replace(',', '').replace('(', '-').replace(')', '')
                    amount = float(amount_str)
                elif 'debit' in col_map or 'credit' in col_map:
                    debit = 0
                    credit = 0
                    if 'debit' in col_map and pd.notna(row[col_map['debit']]):
                        debit = float(str(row[col_map['debit']]).replace('$', '').replace(',', ''))
                    if 'credit' in col_map and pd.notna(row[col_map['credit']]):
                        credit = float(str(row[col_map['credit']]).replace('$', '').replace(',', ''))
                    amount = credit - debit if credit else -debit
                else:
                    continue

                # Get description
                desc = str(row.get(col_map.get('description', ''), 'Unknown'))
                if desc == 'nan':
                    desc = 'Unknown'

                # Get date
                date_val = row.get(col_map.get('trans_date', ''), '')
                if pd.notna(date_val):
                    date_str = str(date_val)
                else:
                    date_str = ''

                transactions.append({
                    'trans_date': date_str,
                    'post_date': str(row.get(col_map.get('post_date', ''), '')) if 'post_date' in col_map else date_str,
                    'description': desc,
                    'amount': amount,
                    'trans_type': 'credit' if amount > 0 else 'debit',
                    'cardholder_name': str(row.get(col_map.get('cardholder', ''), '')) if 'cardholder' in col_map else '',
                    'card_last_four': '',
                    'payment_method': '',
                })
            except (ValueError, TypeError):
                continue

    return transactions, account_info


# --- Word Document Parser ---

def parse_word_document(filepath):
    """Parse a Word document for notes/proof - extracts text and any tables."""
    doc = DocxDocument(filepath)
    content = {
        'text': [],
        'tables': [],
    }

    for para in doc.paragraphs:
        if para.text.strip():
            content['text'].append(para.text.strip())

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content['tables'].append(table_data)

    return content


# --- CSV Bank Statement Parsers ---

def parse_csv_transactions(filepath):
    """Parse CSV bank statements. Auto-detects common bank formats."""
    try:
        df = pd.read_csv(filepath)
    except Exception:
        # Try with different encodings
        try:
            df = pd.read_csv(filepath, encoding='latin-1')
        except Exception:
            return [], {}

    if df.empty:
        return [], {}

    cols = [c.lower().strip() for c in df.columns]

    # Detect bank format by column patterns
    if any('posted date' in c or 'posting date' in c for c in cols):
        return _parse_chase_csv(df)
    elif any('posted' in c and 'date' in c for c in cols) and any('payee' in c for c in cols):
        return _parse_wells_fargo_csv(df)
    elif any('transaction date' in c for c in cols) and any('debit' in c for c in cols):
        return _parse_bofa_csv(df)
    elif any('date' in c for c in cols) and any('reference' in c for c in cols):
        return _parse_citi_csv(df)
    else:
        return _parse_generic_csv(df)


def _parse_chase_csv(df):
    """Parse Chase bank CSV exports."""
    transactions = []
    col_map = {c.lower().strip(): c for c in df.columns}

    date_col = next((col_map[c] for c in col_map if 'posting' in c or 'posted' in c or 'trans' in c and 'date' in c), None)
    desc_col = next((col_map[c] for c in col_map if 'desc' in c or 'memo' in c), None)
    amount_col = next((col_map[c] for c in col_map if 'amount' in c), None)
    type_col = next((col_map[c] for c in col_map if 'type' in c), None)

    if not date_col or not amount_col:
        return _parse_generic_csv(df)

    for _, row in df.iterrows():
        try:
            amount = float(str(row[amount_col]).replace('$', '').replace(',', ''))
            desc = str(row.get(desc_col, 'Unknown')) if desc_col and pd.notna(row.get(desc_col)) else 'Unknown'
            date_str = str(row[date_col])

            transactions.append({
                'trans_date': date_str,
                'post_date': date_str,
                'description': desc,
                'amount': amount,
                'trans_type': str(row.get(type_col, 'debit' if amount < 0 else 'credit')) if type_col else ('debit' if amount < 0 else 'credit'),
                'cardholder_name': '',
                'card_last_four': '',
                'payment_method': '',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Chase', 'account_type': 'bank', 'account_number': ''}


def _parse_wells_fargo_csv(df):
    """Parse Wells Fargo CSV exports."""
    transactions = []
    for _, row in df.iterrows():
        try:
            date_str = ''
            desc = 'Unknown'
            amount = 0

            for col in df.columns:
                cl = col.lower()
                if 'date' in cl:
                    date_str = str(row[col])
                elif 'payee' in cl or 'desc' in cl:
                    if pd.notna(row[col]):
                        desc = str(row[col])
                elif 'amount' in cl:
                    amount = float(str(row[col]).replace('$', '').replace(',', ''))

            transactions.append({
                'trans_date': date_str, 'post_date': date_str, 'description': desc,
                'amount': amount, 'trans_type': 'debit' if amount < 0 else 'credit',
                'cardholder_name': '', 'card_last_four': '', 'payment_method': '',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Wells Fargo', 'account_type': 'bank', 'account_number': ''}


def _parse_bofa_csv(df):
    """Parse Bank of America CSV exports (separate debit/credit columns)."""
    transactions = []
    for _, row in df.iterrows():
        try:
            date_str = ''
            desc = 'Unknown'
            debit = 0
            credit = 0

            for col in df.columns:
                cl = col.lower()
                if 'date' in cl:
                    date_str = str(row[col])
                elif 'desc' in cl or 'payee' in cl:
                    if pd.notna(row[col]):
                        desc = str(row[col])
                elif 'debit' in cl:
                    if pd.notna(row[col]):
                        debit = float(str(row[col]).replace('$', '').replace(',', '').replace('(', '').replace(')', ''))
                elif 'credit' in cl:
                    if pd.notna(row[col]):
                        credit = float(str(row[col]).replace('$', '').replace(',', ''))

            amount = credit - debit if credit else -debit

            transactions.append({
                'trans_date': date_str, 'post_date': date_str, 'description': desc,
                'amount': amount, 'trans_type': 'debit' if amount < 0 else 'credit',
                'cardholder_name': '', 'card_last_four': '', 'payment_method': '',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Bank of America', 'account_type': 'bank', 'account_number': ''}


def _parse_citi_csv(df):
    """Parse Citibank CSV exports."""
    transactions = []
    for _, row in df.iterrows():
        try:
            date_str = ''
            desc = 'Unknown'
            amount = 0

            for col in df.columns:
                cl = col.lower()
                if 'date' in cl and 'status' not in cl:
                    date_str = str(row[col])
                elif 'desc' in cl or 'memo' in cl:
                    if pd.notna(row[col]):
                        desc = str(row[col])
                elif 'amount' in cl or 'debit' in cl or 'credit' in cl:
                    if pd.notna(row[col]):
                        val = str(row[col]).replace('$', '').replace(',', '')
                        if val:
                            amount = float(val)

            transactions.append({
                'trans_date': date_str, 'post_date': date_str, 'description': desc,
                'amount': amount, 'trans_type': 'debit' if amount < 0 else 'credit',
                'cardholder_name': '', 'card_last_four': '', 'payment_method': '',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Citibank', 'account_type': 'bank', 'account_number': ''}


def _parse_generic_csv(df):
    """Parse any CSV with auto-detected columns (fallback)."""
    transactions = []
    col_map = {}

    # Clean column names to prevent KeyErrors from trailing whitespace
    df.columns = [str(c).strip() for c in df.columns]

    for col in df.columns:
        cl = col.lower()
        if 'date' in cl and 'date' not in col_map:
            col_map['date'] = col
        elif 'desc' in cl or 'memo' in cl or 'payee' in cl or 'narrative' in cl:
            col_map['desc'] = col
        elif 'amount' in cl or 'total' in cl:
            col_map['amount'] = col
        elif 'debit' in cl:
            col_map['debit'] = col
        elif 'credit' in cl:
            col_map['credit'] = col
        elif 'card' in cl or 'holder' in cl:
            col_map['cardholder'] = col
        elif 'type' in cl or 'category' in cl:
            col_map['type'] = col

    for _, row in df.iterrows():
        try:
            # Amount
            if 'amount' in col_map:
                val = str(row[col_map['amount']]).replace('$', '').replace(',', '').replace('(', '-').replace(')', '')
                if not val or val == 'nan':
                    continue
                amount = float(val)
            elif 'debit' in col_map or 'credit' in col_map:
                debit = float(str(row.get(col_map.get('debit', ''), 0) or 0).replace('$', '').replace(',', '')) if 'debit' in col_map and pd.notna(row.get(col_map['debit'])) else 0
                credit = float(str(row.get(col_map.get('credit', ''), 0) or 0).replace('$', '').replace(',', '')) if 'credit' in col_map and pd.notna(row.get(col_map['credit'])) else 0
                amount = credit - debit if credit else -debit
            else:
                continue

            desc = str(row.get(col_map.get('desc', ''), 'Unknown'))
            if desc == 'nan':
                desc = 'Unknown'
            date_str = str(row.get(col_map.get('date', ''), ''))
            cardholder = str(row.get(col_map.get('cardholder', ''), '')) if 'cardholder' in col_map else ''
            if cardholder == 'nan':
                cardholder = ''

            transactions.append({
                'trans_date': date_str, 'post_date': date_str, 'description': desc,
                'amount': amount, 'trans_type': 'debit' if amount < 0 else 'credit',
                'cardholder_name': cardholder, 'card_last_four': '', 'payment_method': '',
            })
        except (ValueError, TypeError):
            continue

    return transactions, {'institution': 'Unknown', 'account_type': 'bank', 'account_number': ''}


# --- Master parser dispatcher ---

def parse_document(filepath, doc_type='auto'):
    """
    Parse a document and return transactions + account info.

    doc_type: 'bank_statement', 'credit_card', 'venmo', 'excel', 'word', 'proof', or 'auto'
    """
    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)

    if doc_type == 'auto':
        # Auto-detect based on filename and content
        if ext == '.csv':
            if 'venmo' in filename.lower():
                doc_type = 'venmo'
            else:
                doc_type = 'csv'
        elif ext in ('.xlsx', '.xls'):
            if 'venmo' in filename.lower():
                doc_type = 'venmo'
            else:
                doc_type = 'excel'
        elif ext == '.docx':
            doc_type = 'word'
        elif ext == '.pdf':
            # Try to detect from content
            try:
                pages = parse_pdf_text(filepath)
                full_text = "\n".join(pages).upper()
                if 'CAPITAL ONE' in full_text or 'SPARK CASH' in full_text:
                    doc_type = 'credit_card'
                elif 'VENMO' in full_text and 'STATEMENT' in full_text:
                    doc_type = 'venmo'
                elif 'BANK OF ST' in full_text or 'COMMERCIAL CHECKING' in full_text:
                    doc_type = 'bank_statement'
                else:
                    doc_type = 'bank_statement'  # default for PDFs
            except Exception:
                doc_type = 'bank_statement'

    if doc_type == 'bank_statement':
        return parse_bank_statement(filepath)
    elif doc_type == 'credit_card':
        return parse_capital_one_statement(filepath)
    elif doc_type == 'venmo':
        return parse_venmo_statement(filepath)
    elif doc_type == 'csv':
        return parse_csv_transactions(filepath)
    elif doc_type == 'excel':
        return parse_excel_transactions(filepath)
    elif doc_type in ('word', 'proof'):
        content = parse_word_document(filepath)
        return [], {'content': content, 'doc_type': 'proof'}
    else:
        return [], {}
